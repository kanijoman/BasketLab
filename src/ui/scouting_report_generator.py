"""
Generador de informes de scouting individual en formato DOCX.

Este módulo genera informes detallados de scouting para jugadoras individuales
incluyendo estadísticas, gráficos y espacio para notas del cuerpo técnico.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import Dict, List, Optional, Any
import matplotlib.pyplot as plt
import io
from pathlib import Path
import tempfile
import os
import requests
from PIL import Image
from PyQt6.QtWidgets import QTableWidget, QApplication, QHeaderView, QTableWidgetItem
from PyQt6.QtCore import Qt
from bs4 import BeautifulSoup
from datetime import datetime

from .advanced_stats_calculator import AdvancedStatsCalculator
from .player_stats_table_populator import PlayerStatsTablePopulator
from shotcharts.shot_visualizer import ShotChartVisualizer

from src.utils.collection_utils import is_fbcyl as _is_fbcyl
from shotcharts.zone_analysis import ZoneAnalyzer
from visualization.radar_chart import RadarChart
from .report_document_builder import ReportDocumentBuilder
from .player_data_fetcher import PlayerDataFetcher


class ScoutingReportGenerator:
    """Generador de informes de scouting en formato DOCX."""

    def __init__(self, db_handler: Optional[Any] = None, collection_name: Optional[str] = None):
        """
        Inicializa el generador de informes.

        Args:
            db_handler: Manejador de base de datos para obtener datos adicionales
            collection_name: Nombre de la colección de MongoDB
        """
        self.db_handler = db_handler
        self.collection_name = collection_name
        self.shot_visualizer = ShotChartVisualizer()
        self.zone_analyzer = ZoneAnalyzer()
        self.document_builder = ReportDocumentBuilder()
        self.data_fetcher = PlayerDataFetcher(db_handler, collection_name)

    def generate_team_scouting_report(
        self,
        team_name: str,
        collection_name: str,
        output_path: str,
        player_stats: List[Dict],
        all_player_stats: List[Dict],
        shots_data: Optional[List[Dict]] = None,
        progress_callback: Optional[callable] = None,
        ai_notes: Optional[Dict[str, str]] = None
    ) -> bool:
        """
        Genera un informe de scouting completo para todas las jugadoras de un equipo.

        Args:
            team_name: Nombre del equipo
            collection_name: Nombre de la colección de datos
            output_path: Ruta donde guardar el archivo DOCX
            player_stats: Lista de estadísticas de jugadoras del equipo
            all_player_stats: Lista de estadísticas de todas las jugadoras (para contexto de liga)
            shots_data: Datos de tiros para mapas de calor (opcional)
            progress_callback: Callback para actualizar progreso (opcional)
            ai_notes: Diccionario opcional con notas generadas por IA {player_name: notes_text}

        Returns:
            True si se generó correctamente, False en caso contrario
        """
        try:
            # Crear documento
            doc = Document()

            # Configurar márgenes
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(0.75)
                section.right_margin = Inches(0.75)

            # Añadir título principal
            title = doc.add_heading(f'Informe de Scouting - {team_name}', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Calcular estadísticas avanzadas para todas las jugadoras
            if self.db_handler and collection_name:
                self._calculate_advanced_stats_for_all_players(all_player_stats, collection_name)

            # Filtrar jugadoras del equipo y ordenar por minutos jugados
            team_players = [p for p in player_stats if p.get('team_name') == team_name]
            team_players.sort(key=lambda x: x.get('total_minutes', 0), reverse=True)

            total_players = len(team_players)

            # Generar una página por jugadora
            for idx, player in enumerate(team_players):
                if idx > 0:
                    doc.add_page_break()

                # Actualizar progreso
                if progress_callback:
                    progress_callback(idx + 1, total_players, player.get('player_name', ''))

                self._add_player_page(
                    doc,
                    player,
                    all_player_stats,
                    shots_data,
                    team_name,
                    ai_notes
                )

            # Guardar documento
            doc.save(output_path)

            # Limpiar fotos temporales del caché
            self.data_fetcher.cleanup_photo_cache()

            return True

        except Exception as e:
            import traceback
            traceback.print_exc()

            # Limpiar fotos temporales incluso si hay error
            self.data_fetcher.cleanup_photo_cache()

            return False

    def _add_player_page(
        self,
        doc: Document,
        player: Dict,
        all_player_stats: List[Dict],
        shots_data: Optional[List[Dict]],
        team_name: str,
        ai_notes: Optional[Dict[str, str]] = None
    ):
        """
        Añade una página completa para una jugadora.

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Todas las estadísticas de jugadoras (contexto de liga)
            shots_data: Datos de tiros
            team_name: Nombre del equipo
            ai_notes: Diccionario opcional con notas de IA por jugadora
        """
        # Sección de datos generales con foto
        self._add_player_header(doc, player, team_name)

        # Sección 1: Estadística General
        self._add_basic_stats_section(doc, player, all_player_stats)

        # Sección 2: Estadística Avanzada
        self._add_advanced_stats_section(doc, player, all_player_stats)

        # Sección 3: Perfil de Lanzamiento
        self._add_shooting_profile_section(doc, player, shots_data, team_name)

        # Sección 4: Perfil de Juego
        self._add_game_profile_section(doc, player, all_player_stats)

        # Sección 5: Notas del cuerpo técnico (con notas de IA si están disponibles)
        player_name = player.get('player_name', '')
        player_ai_notes = ai_notes.get(player_name) if ai_notes else None
        self._add_notes_section(doc, player_ai_notes)

    def _add_player_header(self, doc: Document, player: Dict, team_name: str):
        """
        Añade el encabezado con datos generales y foto de la jugadora.

        Args:
            doc: Documento DOCX
            player: Diccionario con datos de la jugadora
            team_name: Nombre del equipo
        """
        # Crear tabla de 2 columnas para datos + foto
        table = doc.add_table(rows=1, cols=2)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Columna izquierda: Datos generales
        left_cell = table.rows[0].cells[0]
        left_cell.width = Inches(4.5)

        # Obtener datos de la jugadora (algunos pueden no estar disponibles)
        player_name = player.get('player_name', '')
        player_id = player.get('player_id', '')

        # Obtener dorsal, URL de foto, y team_id desde el último partido
        dorsal, photo_url, team_id = self.data_fetcher.get_player_dorsal_and_photo(player_id)

        # Obtener fecha de nacimiento, edad y altura desde FEB
        birth_date, age, height = self.data_fetcher.get_player_birth_info(player_id, team_id)

        # Posición (no disponible, dejar en blanco)
        position = ""

        # Añadir información - usar el párrafo existente de la celda
        p = left_cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        self.document_builder.add_formatted_text(p, f"#{dorsal}# ", bold=True, size=16)
        self.document_builder.add_formatted_text(p, player_name, bold=True, size=16)

        # Fecha de nacimiento
        p = left_cell.add_paragraph()
        if birth_date:
            self.document_builder.add_formatted_text(p, f"Fecha de nacimiento: {birth_date}", size=11)
        else:
            self.document_builder.add_formatted_text(p, "Fecha de nacimiento:", size=11)

        # Edad
        p = left_cell.add_paragraph()
        if age:
            self.document_builder.add_formatted_text(p, f"Edad: {age} años", size=11)
        else:
            self.document_builder.add_formatted_text(p, "Edad:", size=11)

        # Posición
        p = left_cell.add_paragraph()
        if position:
            self.document_builder.add_formatted_text(p, f"Posición: {position}", size=11)
        else:
            self.document_builder.add_formatted_text(p, "Posición:", size=11)

        # Altura
        p = left_cell.add_paragraph()
        if height:
            self.document_builder.add_formatted_text(p, f"Altura: {height}", size=11)
        else:
            self.document_builder.add_formatted_text(p, "Altura:", size=11)

        # Columna derecha: Foto de la jugadora
        right_cell = table.rows[0].cells[1]
        right_cell.width = Inches(1.5)
        right_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Intentar descargar y añadir foto de la jugadora
        photo_added = False

        if photo_url:
            photo_path = self.data_fetcher.download_photo_from_url(photo_url, player_id)
            if photo_path and os.path.exists(photo_path):
                try:
                    p = right_cell.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = p.add_run()
                    run.add_picture(photo_path, width=Inches(1.2))
                    photo_added = True
                except Exception as e:
                    pass

        # Si no se pudo añadir foto, mostrar placeholder
        if not photo_added:
            p = right_cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.text = "[FOTO]"
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(150, 150, 150)

        # Añadir borde al placeholder de foto
        self.document_builder.set_cell_border(right_cell)

    def _add_basic_stats_section(self, doc: Document, player: Dict, all_player_stats: List[Dict]):
        """
        Añade la sección de estadísticas básicas usando imágenes PNG.
        Genera dos imágenes: una para promedios y otra para totales.

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Lista de todas las jugadoras para calcular cuartiles
        """
        # Título de sección
        heading = doc.add_heading('1. Estadística General', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Subtítulo: Promedios
        p = doc.add_paragraph()
        run = p.add_run('Promedios por partido:')
        run.font.bold = True
        run.font.size = Pt(10)

        # Generar imagen de estadísticas básicas (promedios)
        stats_image_avg = self._generate_stats_image_single_row(player, all_player_stats, view_mode='average')

        if stats_image_avg and os.path.exists(stats_image_avg):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(stats_image_avg, width=Inches(6.0))
            except Exception as e:
                pass

        # Subtítulo: Totales
        p = doc.add_paragraph()
        run = p.add_run('Totales acumulados:')
        run.font.bold = True
        run.font.size = Pt(10)

        # Generar imagen de estadísticas básicas (totales)
        stats_image_total = self._generate_stats_image_single_row(player, all_player_stats, view_mode='total')

        if stats_image_total and os.path.exists(stats_image_total):
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(stats_image_total, width=Inches(6.0))
            except Exception as e:
                pass

        # Si ambas fallan, usar tabla simple como fallback
        if not (stats_image_avg and stats_image_total):
            self._add_basic_stats_table_fallback(doc, player)

    def _add_basic_stats_table_fallback(self, doc: Document, player: Dict):
        """
        Añade tabla de estadísticas básicas (fallback si falla la imagen).

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
        """
        # Obtener datos
        games = player.get('games_played', 0)
        total_mins = player.get('total_minutes', 0)
        mins_per_game = total_mins / games if games > 0 else 0

        # Obtener datos
        games = player.get('games_played', 0)
        total_mins = player.get('total_minutes', 0)
        mins_per_game = total_mins / games if games > 0 else 0

        # Crear tabla de estadísticas
        stats_table = doc.add_table(rows=3, cols=8)
        stats_table.style = 'Light Grid Accent 1'

        # Encabezados
        headers = ['PJ', 'Min', 'Pts', 'Reb', 'RO', 'RD', 'Ast', 'Rob']
        for i, header in enumerate(headers):
            cell = stats_table.rows[0].cells[i]
            cell.text = header
            self.document_builder.format_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Totales
        totals_row = stats_table.rows[1]
        totals_data = [
            str(games),
            f"{total_mins // 60}:{total_mins % 60:02d}",
            str(player.get('total_pts', 0)),
            str(player.get('total_rt', 0)),
            str(player.get('total_ro', 0)),
            str(player.get('total_rd', 0)),
            str(player.get('total_assist', 0)),
            str(player.get('total_st', 0))
        ]
        for i, data in enumerate(totals_data):
            totals_row.cells[i].text = data
            self.document_builder.format_cell(totals_row.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

        # Promedios
        avg_row = stats_table.rows[2]
        ppg = player.get('total_pts', 0) / games if games > 0 else 0
        rpg = player.get('total_rt', 0) / games if games > 0 else 0
        ropg = player.get('total_ro', 0) / games if games > 0 else 0
        rdpg = player.get('total_rd', 0) / games if games > 0 else 0
        apg = player.get('total_assist', 0) / games if games > 0 else 0
        spg = player.get('total_st', 0) / games if games > 0 else 0

        avg_data = [
            '',  # PJ no tiene promedio
            f"{mins_per_game:.1f}",
            f"{ppg:.1f}",
            f"{rpg:.1f}",
            f"{ropg:.1f}",
            f"{rdpg:.1f}",
            f"{apg:.1f}",
            f"{spg:.1f}"
        ]
        for i, data in enumerate(avg_data):
            avg_row.cells[i].text = data
            self.document_builder.format_cell(avg_row.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

        # Segunda tabla: Más estadísticas
        doc.add_paragraph()  # Espaciado
        stats_table2 = doc.add_table(rows=3, cols=8)
        stats_table2.style = 'Light Grid Accent 1'

        # Encabezados
        headers2 = ['BP', 'Tap', 'FP', 'FR', '+/-', 'Val', '%T2', '%T3']
        for i, header in enumerate(headers2):
            cell = stats_table2.rows[0].cells[i]
            cell.text = header
            self.document_builder.format_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Totales
        totals_row2 = stats_table2.rows[1]
        totals_data2 = [
            str(player.get('total_to', 0)),
            str(player.get('total_bs', 0)),
            str(player.get('total_pf', 0)),
            str(player.get('total_rf', 0)),
            str(player.get('total_pllss', 0)),
            str(player.get('total_val', 0)),
            f"{player.get('fg2_percentage', 0):.1f}%",
            f"{player.get('fg3_percentage', 0):.1f}%"
        ]
        for i, data in enumerate(totals_data2):
            totals_row2.cells[i].text = data
            self.document_builder.format_cell(totals_row2.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

        # Promedios
        avg_row2 = stats_table2.rows[2]
        topg = player.get('total_to', 0) / games if games > 0 else 0
        bpg = player.get('total_bs', 0) / games if games > 0 else 0
        fpg = player.get('total_pf', 0) / games if games > 0 else 0
        frpg = player.get('total_rf', 0) / games if games > 0 else 0
        pllss_pg = player.get('total_pllss', 0) / games if games > 0 else 0
        val_pg = player.get('total_val', 0) / games if games > 0 else 0

        avg_data2 = [
            f"{topg:.1f}",
            f"{bpg:.1f}",
            f"{fpg:.1f}",
            f"{frpg:.1f}",
            f"{pllss_pg:.1f}",
            f"{val_pg:.1f}",
            '',  # Los porcentajes no se promedian de nuevo
            ''
        ]
        for i, data in enumerate(avg_data2):
            avg_row2.cells[i].text = data
            self.document_builder.format_cell(avg_row2.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

    def _add_advanced_stats_section(self, doc: Document, player: Dict, all_player_stats: List[Dict]):
        """
        Añade la sección de estadísticas avanzadas usando imagen PNG.

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Lista de todas las jugadoras para calcular cuartiles
        """
        # Título de sección
        heading = doc.add_heading('2. Estadística Avanzada', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Generar imagen de estadísticas avanzadas (promedios)
        stats_image = self._generate_advanced_stats_image(player, all_player_stats)

        if stats_image and os.path.exists(stats_image):
            try:
                # Añadir imagen al documento
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(stats_image, width=Inches(6.0))
            except Exception as e:
                # Si falla, usar tabla simple como fallback
                pass
                self._add_advanced_stats_table_fallback(doc, player)
        else:
            # Fallback a tabla simple
            self._add_advanced_stats_table_fallback(doc, player)

    def _add_advanced_stats_table_fallback(self, doc: Document, player: Dict):
        """
        Añade tabla de estadísticas avanzadas (fallback si falla la imagen).

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
        """
        # Obtener datos de estadísticas avanzadas
        games = player.get('games_played', 0)

        # Crear tabla de estadísticas avanzadas
        stats_table = doc.add_table(rows=2, cols=9)
        stats_table.style = 'Light Grid Accent 1'

        # Encabezados
        headers = ['TS%', 'eFG%', 'FTr', '3Pr', 'ORtg', 'DRtg', 'USG%', '%AST', '%TO']
        for i, header in enumerate(headers):
            cell = stats_table.rows[0].cells[i]
            cell.text = header
            self.document_builder.format_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Valores
        values_row = stats_table.rows[1]
        values = [
            f"{player.get('ts', 0):.1f}%",
            f"{player.get('efg', 0):.1f}%",
            f"{player.get('ftr', 0):.1f}%",
            f"{player.get('three_pr', 0):.1f}%",
            f"{player.get('orating', 0):.1f}",
            f"{player.get('drating', 0):.1f}",
            f"{player.get('usage', 0):.1f}%",
            f"{player.get('ast_pct', 0):.1f}%",
            f"{player.get('tov_pct', 0):.1f}%"
        ]
        for i, value in enumerate(values):
            values_row.cells[i].text = value
            self.document_builder.format_cell(values_row.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

        # Segunda fila de estadísticas avanzadas
        doc.add_paragraph()  # Espaciado
        stats_table2 = doc.add_table(rows=2, cols=6)
        stats_table2.style = 'Light Grid Accent 1'

        # Encabezados
        headers2 = ['%ROB', '%TAP', '%RO', '%RD', 'AST Ratio', 'AST/TO']
        for i, header in enumerate(headers2):
            cell = stats_table2.rows[0].cells[i]
            cell.text = header
            self.document_builder.format_cell(cell, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        # Valores
        values_row2 = stats_table2.rows[1]
        values2 = [
            f"{player.get('stl_pct', 0):.1f}%",
            f"{player.get('blk_pct', 0):.1f}%",
            f"{player.get('orb_pct', 0):.1f}%",
            f"{player.get('drb_pct', 0):.1f}%",
            f"{player.get('ast_ratio', 0):.1f}",
            f"{player.get('ast_to_ratio', 0):.2f}"
        ]
        for i, value in enumerate(values2):
            values_row2.cells[i].text = value
            self.document_builder.format_cell(values_row2.cells[i], align=WD_ALIGN_PARAGRAPH.CENTER)

    def _add_shooting_profile_section(
        self,
        doc: Document,
        player: Dict,
        shots_data: Optional[List[Dict]],
        team_name: str
    ):
        """
        Añade la sección de perfil de lanzamiento con mapa de calor y gráfico por zonas.

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
            shots_data: Datos de tiros
            team_name: Nombre del equipo
        """
        # Título de sección
        heading = doc.add_heading('3. Perfil de Lanzamiento', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        if not shots_data:
            p = doc.add_paragraph()
            p.add_run("No hay datos de tiros disponibles.").italic = True
            return

        try:
            # Obtener datos de la jugadora
            player_id = player.get('player_id', '')
            player_uuid = player.get('player_uuid', '')
            player_name = player.get('player_name', '')

            # Detectar si es FEB o FBCYL basándonos en el collection_name
            is_fbcyl = _is_fbcyl(self.collection_name)

            # Filtrar tiros según la competición
            player_shots = []

            if is_fbcyl:
                # FBCYL: Usar UUID (preferido) o normalized_name (fallback para casos sin UUID)
                # Calcular normalized_name para fallback
                normalized_name = ""
                if player_name:
                    words = player_name.strip().split()
                    if len(words) >= 2:
                        initial = words[0][0] if words[0] else ""
                        surnames = words[-2:] if len(words) >= 2 else words[-1:]
                        normalized_name = f"{initial} {' '.join(surnames)}"
                    else:
                        normalized_name = player_name

                for shot in shots_data:
                    # Intentar match por UUID (más preciso)
                    if player_uuid and shot.get('player_uuid'):
                        if shot.get('player_uuid') == player_uuid:
                            player_shots.append(shot)
                            continue

                    # Fallback: normalized_name (cuando UUID no disponible)
                    if normalized_name and shot.get('normalized_name'):
                        if shot.get('normalized_name') == normalized_name:
                            player_shots.append(shot)
                            continue
            else:
                # FEB: Usar player_id (suficiente y más preciso)
                for shot in shots_data:
                    if player_id and shot.get('player_id'):
                        if shot.get('player_id') == player_id:
                            player_shots.append(shot)

            if not player_shots:
                p = doc.add_paragraph()
                p.add_run("No hay datos de tiros para esta jugadora.").italic = True
                return

            # Crear tabla para los dos gráficos
            table = doc.add_table(rows=1, cols=2)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Mapa de calor
            heatmap_path = self._generate_heatmap(player_shots, player.get('player_name', ''))
            if heatmap_path and os.path.exists(heatmap_path):
                left_cell = table.rows[0].cells[0]
                p = left_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(heatmap_path, width=Inches(2.8))
                # Limpiar archivo temporal
                try:
                    os.unlink(heatmap_path)
                except:
                    pass
            else:
                # Si no hay mapa de calor, añadir texto explicativo
                left_cell = table.rows[0].cells[0]
                p = left_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run("No hay suficientes aciertos\npara trazar un mapa de calor")
                run.italic = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(100, 100, 100)

            # Gráfico por zonas
            zones_path = self._generate_zone_chart(player_shots, player.get('player_name', ''))
            if zones_path and os.path.exists(zones_path):
                right_cell = table.rows[0].cells[1]
                p = right_cell.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(zones_path, width=Inches(2.8))
                # Limpiar archivo temporal
                try:
                    os.unlink(zones_path)
                except:
                    pass

        except Exception as e:
            p = doc.add_paragraph()
            p.add_run("Error al generar gráficos de lanzamiento.").italic = True

    def _add_game_profile_section(
        self,
        doc: Document,
        player: Dict,
        all_player_stats: List[Dict]
    ):
        """
        Añade la sección de perfil de juego con radar chart.

        Args:
            doc: Documento DOCX
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Todas las estadísticas de jugadoras (para contexto)
        """
        # Título de sección
        heading = doc.add_heading('4. Perfil de Juego', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT

        try:
            # Generar radar chart
            radar_path = self._generate_radar_chart(player, all_player_stats)

            if radar_path and os.path.exists(radar_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(radar_path, width=Inches(4.5))

                # Limpiar archivo temporal
                try:
                    os.unlink(radar_path)
                except:
                    pass
            else:
                p = doc.add_paragraph()
                p.add_run("No se pudo generar el radar chart.").italic = True

        except Exception as e:
            import traceback
            traceback.print_exc()
            p = doc.add_paragraph()
            p.add_run("Error al generar radar chart.").italic = True

    def _add_notes_section(self, doc: Document, ai_notes: Optional[str] = None):
        """
        Añade la sección de notas para el cuerpo técnico.

        Args:
            doc: Documento DOCX
            ai_notes: Notas generadas por IA (opcional)
        """
        # Título de sección
        heading = doc.add_heading('5. Notas del Cuerpo Técnico', level=2)
        heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        # Reducir espaciado después del título
        heading.paragraph_format.space_after = Pt(6)

        # Si hay notas de IA, añadirlas
        if ai_notes:
            # Procesar las notas para mantener formato con bullets
            for line in ai_notes.split('\n'):
                line = line.strip()
                if not line:
                    continue

                # Detectar si es un encabezado (FORTALEZAS, DEBILIDADES, PERFIL)
                if line.startswith('**') and line.endswith('**'):
                    # Es un encabezado
                    p = doc.add_paragraph()
                    run = p.add_run(line.replace('**', ''))
                    run.font.bold = True
                    run.font.size = Pt(11)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(3)
                elif line.startswith('•') or line.startswith('-'):
                    # Es un bullet point
                    text = line.lstrip('•-').strip()
                    p = doc.add_paragraph(text, style='List Bullet')
                    p.paragraph_format.space_after = Pt(3)
                    # Formatear el texto
                    for run in p.runs:
                        run.font.size = Pt(10)
                else:
                    # Es texto normal
                    p = doc.add_paragraph(line)
                    p.paragraph_format.space_after = Pt(3)
                    for run in p.runs:
                        run.font.size = Pt(10)
        else:
            # Añadir texto placeholder para notas manuales
            p = doc.add_paragraph()
            run = p.add_run("[Espacio para observaciones y notas del cuerpo técnico]")
            run.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)
            # Sin espaciado después
            p.paragraph_format.space_after = Pt(0)

    # Métodos auxiliares

    def _calculate_advanced_stats_for_all_players(self, player_stats: List[Dict], collection_name: str):
        """
        Calcula estadísticas avanzadas para todas las jugadoras.

        Args:
            player_stats: Lista de estadísticas de jugadoras
            collection_name: Nombre de la colección
        """
        # Group players by team to get team and opponent stats
        teams = {}
        for player in player_stats:
            team_name = player['team_name']
            if team_name not in teams:
                teams[team_name] = {
                    'team_stats': self.db_handler.get_aggregated_team_stats(
                        collection_name, team_name
                    ),
                    'opp_stats': self.db_handler.get_aggregated_opponent_stats(
                        collection_name, team_name
                    )
                }

        # Calculate advanced stats for each player
        calculator = AdvancedStatsCalculator()
        for player in player_stats:
            team_name = player['team_name']
            team_data = teams.get(team_name, {})
            team_stats = team_data.get('team_stats', {})
            opp_stats = team_data.get('opp_stats', {})

            # Calculate shooting percentages first
            fg2_made = player.get('total_p2m', 0)
            fg2_attempted = player.get('total_p2a', 0)
            fg3_made = player.get('total_p3m', 0)
            fg3_attempted = player.get('total_p3a', 0)
            ft_made = player.get('total_p1m', 0)
            ft_attempted = player.get('total_p1a', 0)
            games = player.get('games_played', 0)

            player['fg2_pct'] = (fg2_made / fg2_attempted * 100) if fg2_attempted > 0 else 0
            player['fg3_pct'] = (fg3_made / fg3_attempted * 100) if fg3_attempted > 0 else 0
            player['ft_pct'] = (ft_made / ft_attempted * 100) if ft_attempted > 0 else 0

            # Calculate per-game stats
            # Guardar minutes_per_game en segundos primero para cálculos
            player['minutes_per_game'] = (player.get('total_minutes', 0) / games) if games > 0 else 0
            # mpg es la versión en minutos (dividido por 60) para las estadísticas avanzadas
            player['mpg'] = (player['minutes_per_game'] / 60) if player['minutes_per_game'] > 0 else 0
            player['ppg'] = player.get('total_pts', 0) / games if games > 0 else 0
            player['val_pg'] = player.get('total_val', 0) / games if games > 0 else 0

            # Calculate advanced stats that don't need team data
            fga = fg2_attempted + fg3_attempted
            player['efg'] = calculator.calculate_effective_fg_percentage(player)
            player['ts'] = calculator.calculate_true_shooting_percentage(player)
            player['ftr'] = calculator.calculate_ftr(player)
            player['three_pr'] = calculator.calculate_3pr(player)

            # Calculate rebound percentages (simple version without team data)
            # Nota: Estos porcentajes son estimaciones sin datos de equipo/oponente
            # Estimación de rebotes disponibles: asumimos ~80 rebotes por partido (40 por equipo)
            estimated_team_rebounds = games * 40 if games > 0 else 1
            player['orb_pct'] = (player.get('total_ro', 0) / estimated_team_rebounds * 100) if estimated_team_rebounds > 0 else 0.0
            player['drb_pct'] = (player.get('total_rd', 0) / estimated_team_rebounds * 100) if estimated_team_rebounds > 0 else 0.0

            # Calculate other percentages based on possessions estimate
            possessions_est = fga + 0.44 * ft_attempted + player.get('total_to', 0)
            player['ast_pct'] = (player.get('total_as', player.get('total_assist', 0)) / possessions_est * 100) if possessions_est > 0 else 0.0
            player['tov_pct'] = (player.get('total_to', 0) / possessions_est * 100) if possessions_est > 0 else 0.0
            player['stl_pct'] = (player.get('total_st', 0) / possessions_est * 100) if possessions_est > 0 else 0.0
            player['blk_pct'] = (player.get('total_bl', player.get('total_bs', 0)) / possessions_est * 100) if possessions_est > 0 else 0.0

            # Calculate team-dependent stats if data is available
            if team_stats and opp_stats:
                # Override with more accurate calculations using team data
                full_advanced_stats = calculator.calculate_all_advanced_stats(player, team_stats, opp_stats)
                player.update(full_advanced_stats)
            else:
                # Approximations for team-dependent stats
                player['usage'] = (possessions_est / games) if games > 0 else 0.0
                player['orating'] = (player.get('total_pts', 0) / possessions_est * 100) if possessions_est > 0 else 0.0
                # Defensive rating approximation (lower is better, league average ~110)
                defensive_contrib = player.get('total_st', 0) + player.get('total_bl', player.get('total_bs', 0)) + player.get('total_rd', 0)
                player['drating'] = max(80, 120 - (defensive_contrib / games * 2)) if games > 0 else 110.0

            # Calculate additional ratios
            total_ast = player.get('total_assist', player.get('total_as', 0))
            total_to = player.get('total_to', 0)
            total_fga = fga
            total_fta = ft_attempted

            possessions = total_fga + (0.44 * total_fta) + total_ast + total_to
            player['ast_ratio'] = (100 * total_ast / possessions) if possessions > 0 else 0
            player['ast_to_ratio'] = total_ast / total_to if total_to > 0 else 0

            ast_pct = player.get('ast_pct', 0)
            usage = player.get('usage', 0)
            player['ast_usg'] = (ast_pct / usage * 100) if usage > 0 else 0

    def _generate_stats_image_single_row(
        self,
        player: Dict,
        all_player_stats: List[Dict],
        view_mode: str
    ) -> Optional[str]:
        """
        Genera una imagen PNG de una fila de estadísticas básicas con coloreado por cuartiles.

        Args:
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Lista de todas las jugadoras para calcular cuartiles correctos
            view_mode: Modo de vista ('total' o 'average')

        Returns:
            Ruta al archivo temporal con la imagen, o None si falla
        """
        try:
            # Crear QApplication si no existe
            app = QApplication.instance()
            if app is None:
                app = QApplication([])

            # Definir columnas de estadísticas (sin Jugadora ni Equipo para informe individual)
            PLAYER_COLUMNS = [
                "PJ", "Min", "Pts", "TL%", "T2%", "T3%",
                "RO", "RD", "RT", "Ast", "Rec", "BP", "Tap", "FP", "FR",
                "+/-", "Val"
            ]

            # Crear tabla con configuración completa - 1 fila
            table = QTableWidget()
            table.setRowCount(1)
            table.setColumnCount(len(PLAYER_COLUMNS))
            table.setHorizontalHeaderLabels(PLAYER_COLUMNS)

            # Configurar tabla (igual que en PlayerStatsWindow)
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
            table.horizontalHeader().setStretchLastSection(False)
            table.verticalHeader().setVisible(False)
            table.setAlternatingRowColors(True)
            table.setSortingEnabled(False)

            # Ajustar anchos de columna para caber en página (reducidos agresivamente)
            for i in range(len(PLAYER_COLUMNS)):
                table.horizontalHeader().resizeSection(i, 42)  # Ancho uniforme compacto

            # Crear lista con un solo jugador
            player_stats = [player]

            # Calcular cuartiles manualmente usando nombres de campos
            from .stats_config import get_quartile_color, calculate_quartiles as calc_quartiles
            from .table_items import NumericTableWidgetItem
            from .player_stats_calculator import PlayerStatsCalculator
            
            quartiles = {}
            
            # Definir los campos para los que necesitamos cuartiles
            if view_mode == 'average':
                stat_fields = {
                    'minutes_per_game': 'total_minutes',
                    'pts': 'total_pts',
                    'tl_pct': 'tl_pct', 
                    't2_pct': 't2_pct',
                    't3_pct': 't3_pct',
                    'ro': 'total_ro',
                    'rd': 'total_rd',
                    'rt': 'total_rt',
                    'assist': 'total_assist',
                    'st': 'total_st',
                    'to': 'total_to',
                    'bs': 'total_bs',
                    'pf': 'total_pf',
                    'rf': 'total_rf',
                    'pllss': 'total_pllss',
                    'val': 'total_val'
                }
            else:  # total
                stat_fields = {
                    'pts': 'total_pts',
                    'tl_pct': 'tl_pct',
                    't2_pct': 't2_pct',
                    't3_pct': 't3_pct',
                    'ro': 'total_ro',
                    'rd': 'total_rd',
                    'rt': 'total_rt',
                    'assist': 'total_assist',
                    'st': 'total_st',
                    'to': 'total_to',
                    'bs': 'total_bs',
                    'pf': 'total_pf',
                    'rf': 'total_rf',
                    'pllss': 'total_pllss',
                    'val': 'total_val'
                }
            
            # Calcular cuartiles para cada campo
            for field_key, total_key in stat_fields.items():
                values = []
                for p in all_player_stats:
                    games = p.get('games_played', 0)
                    if games == 0:
                        continue
                    
                    # Calcular el valor según el campo
                    if field_key == 'tl_pct':
                        val = p.get('total_p1m', 0) / p.get('total_p1a', 1) * 100 if p.get('total_p1a', 0) > 0 else 0
                    elif field_key == 't2_pct':
                        val = p.get('total_p2m', 0) / p.get('total_p2a', 1) * 100 if p.get('total_p2a', 0) > 0 else 0
                    elif field_key == 't3_pct':
                        val = p.get('total_p3m', 0) / p.get('total_p3a', 1) * 100 if p.get('total_p3a', 0) > 0 else 0
                    elif field_key == 'minutes_per_game':
                        # Convertir de segundos a minutos
                        val = p.get(total_key, 0) / 60.0 / games if games > 0 else 0
                    else:
                        val = p.get(total_key, 0)
                        if view_mode == 'average':
                            val = val / games if games > 0 else 0
                    
                    if val > 0 or field_key in ['pllss']:  # +/- puede ser negativo
                        values.append(val)
                
                if len(values) >= 4:
                    quartiles[field_key] = calc_quartiles(values)

            row = 0
            games = player.get('games_played', 0)

            if view_mode == 'average':
                # Columna 0: PJ
                table.setItem(row, 0, NumericTableWidgetItem(games, str(games)))

                # Columna 1: Min - convertir de segundos a minutos
                mins_per_game_seconds = player.get('minutes_per_game', 0)
                mins_per_game = mins_per_game_seconds / 60.0 if mins_per_game_seconds else 0
                item = NumericTableWidgetItem(mins_per_game, f"{mins_per_game:.1f}")
                # Aplicar coloreado por cuartiles para minutos
                if 'minutes_per_game' in quartiles:
                    item.setBackground(get_quartile_color(mins_per_game, quartiles['minutes_per_game'], False))
                table.setItem(row, 1, item)

                # Columna 2: Pts
                ppg = player.get('total_pts', 0) / games if games > 0 else 0
                item = NumericTableWidgetItem(ppg, f"{ppg:.1f}")
                if 'pts' in quartiles:
                    item.setBackground(get_quartile_color(ppg, quartiles['pts'], False))
                table.setItem(row, 2, item)

                # Columnas 3-5: Porcentajes de tiro
                ft_pct = player.get('total_p1m', 0) / player.get('total_p1a', 1) * 100 if player.get('total_p1a', 0) > 0 else 0
                fg2_pct = player.get('total_p2m', 0) / player.get('total_p2a', 1) * 100 if player.get('total_p2a', 0) > 0 else 0
                fg3_pct = player.get('total_p3m', 0) / player.get('total_p3a', 1) * 100 if player.get('total_p3a', 0) > 0 else 0

                for idx, (val, key) in enumerate([(ft_pct, 'tl_pct'), (fg2_pct, 't2_pct'), (fg3_pct, 't3_pct')], start=3):
                    item = NumericTableWidgetItem(val, f"{val:.1f}%")
                    if key in quartiles:
                        item.setBackground(get_quartile_color(val, quartiles[key], False))
                    table.setItem(row, idx, item)

                # Resto de estadísticas por juego
                stat_keys = ['ro', 'rd', 'rt', 'assist', 'st', 'to', 'bs', 'pf', 'rf', 'pllss', 'val']
                stat_totals = ['total_ro', 'total_rd', 'total_rt', 'total_assist', 'total_st', 'total_to', 
                              'total_bs', 'total_pf', 'total_rf', 'total_pllss', 'total_val']

                for col_idx, (key, total_key) in enumerate(zip(stat_keys, stat_totals), start=6):
                    val = player.get(total_key, 0) / games if games > 0 else 0
                    item = NumericTableWidgetItem(val, f"{val:.1f}")
                    if key in quartiles:
                        reverse = key in ['to', 'pf']  # Menos es mejor
                        item.setBackground(get_quartile_color(val, quartiles[key], reverse))
                    table.setItem(row, col_idx, item)

            else:  # total
                # Similar pero con totales
                table.setItem(row, 0, NumericTableWidgetItem(games, str(games)))

                # Columna 1: Min - convertir segundos a formato MM:SS
                total_seconds = player.get('total_minutes', 0)
                total_mins = int(total_seconds // 60)
                secs = int(total_seconds % 60)
                mins_display = f"{total_mins}:{secs:02d}"
                table.setItem(row, 1, NumericTableWidgetItem(total_seconds, mins_display))

                # Totales directos
                stat_keys = ['pts', 'tl_pct', 't2_pct', 't3_pct', 'ro', 'rd', 'rt', 'assist', 'st', 'to', 'bs', 'pf', 'rf', 'pllss', 'val']
                stat_totals = ['total_pts', 'total_p1m', 'total_p2m', 'total_p3m', 'total_ro', 'total_rd', 'total_rt',
                              'total_assist', 'total_st', 'total_to', 'total_bs', 'total_pf', 'total_rf', 'total_pllss', 'total_val']

                for col_idx, (key, total_key) in enumerate(zip(stat_keys, stat_totals), start=2):
                    val = player.get(total_key, 0)
                    
                    # Lanzamientos: mostrar formato acierto-fallo en totales
                    if key == 'tl_pct':
                        made = int(player.get('total_p1m', 0))
                        attempted = int(player.get('total_p1a', 0))
                        item = NumericTableWidgetItem(made, f"{made}-{attempted}")
                    elif key == 't2_pct':
                        made = int(player.get('total_p2m', 0))
                        attempted = int(player.get('total_p2a', 0))
                        item = NumericTableWidgetItem(made, f"{made}-{attempted}")
                    elif key == 't3_pct':
                        made = int(player.get('total_p3m', 0))
                        attempted = int(player.get('total_p3a', 0))
                        item = NumericTableWidgetItem(made, f"{made}-{attempted}")
                    else:
                        item = NumericTableWidgetItem(val, str(int(val)))
                    
                    if key in quartiles:
                        reverse = key in ['to', 'pf']
                        item.setBackground(get_quartile_color(val, quartiles[key], reverse))
                    table.setItem(row, col_idx, item)

            # Solo ajustar altura de filas, mantener anchos fijos
            table.resizeRowsToContents()

            # Calcular tamaño total necesario con margen de seguridad
            total_width = sum(table.columnWidth(i) for i in range(table.columnCount())) + 4
            total_height = (sum(table.rowHeight(i) for i in range(table.rowCount())) +
                          table.horizontalHeader().height() + 4)

            # Establecer tamaño mínimo (no fijo) para permitir expansión si es necesario
            table.setMinimumSize(total_width, total_height)
            table.resize(total_width, total_height)

            # Evitar que la ventana se muestre en pantalla (prevenir parpadeo)
            table.setAttribute(Qt.WidgetAttribute.WA_DontShowOnScreen, True)

            # Asegurar que la tabla se renderice correctamente
            table.show()
            app.processEvents()

            # Forzar otro ajuste de tamaño después de mostrar
            app.processEvents()

            # Capturar la tabla como imagen usando el rect completo
            pixmap = table.grab(table.rect())

            # Ocultar la tabla
            table.hide()
            table.deleteLater()

            # Guardar en archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_file.close()

            if pixmap.save(temp_file.name, "PNG", quality=100):
                return temp_file.name
            else:
                return None

        except Exception as e:
            import traceback
            traceback.print_exc()
            return None

    def _generate_advanced_stats_image(self, player: Dict, all_player_stats: List[Dict]) -> Optional[str]:
        """
        Genera una imagen PNG de las estadísticas avanzadas usando el sistema de tablas existente.

        Args:
            player: Diccionario con estadísticas de la jugadora
            all_player_stats: Lista de todas las jugadoras para calcular cuartiles correctos

        Returns:
            Ruta al archivo temporal con la imagen, o None si falla
        """
        try:
            # Crear QApplication si no existe
            app = QApplication.instance()
            if app is None:
                app = QApplication([])

            # Columnas de estadísticas avanzadas (sin Jugadora ni Equipo para informe individual)
            ADVANCED_COLUMNS = [
                "PJ", "Min/PJ", "Pts/PJ", "TS%", "eFG%",
                "3PAr", "FTr", "ORB%", "DRB%", "AST%", "TO%",
                "STL%", "BLK%", "USG%", "ORtg", "DRtg", "Val/PJ"
            ]

            # Crear tabla
            table = QTableWidget()
            table.setColumnCount(len(ADVANCED_COLUMNS))
            table.setHorizontalHeaderLabels(ADVANCED_COLUMNS)

            # Configurar tabla
            table.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)
            table.horizontalHeader().setStretchLastSection(False)
            table.verticalHeader().setVisible(False)
            table.setAlternatingRowColors(True)
            table.setSortingEnabled(False)

            # Ajustar anchos de columna (compactos para caber en página)
            for i in range(len(ADVANCED_COLUMNS)):
                table.horizontalHeader().resizeSection(i, 42)  # Ancho uniforme compacto

            # Crear lista con un jugador
            player_stats = [player]

            # Poblar usando el método de ventana avanzada
            table.setRowCount(1)
            row = 0

            # Calcular cuartiles para estadísticas avanzadas
            from .stats_config import get_quartile_color, calculate_quartiles

            # Definición de campos avanzados con sus configuraciones (sin columnas 0 y 1)
            # NOTA: Los nombres deben coincidir con los que devuelve AdvancedStatsCalculator
            ADVANCED_STAT_FIELDS = {
                0: ('games_played', False),  # PJ
                1: ('minutes_per_game', False),  # Min/PJ
                2: ('ppg', False),
                3: ('ts', False),           # TS%
                4: ('efg', False),          # eFG%
                5: ('three_pr', False),     # 3PAr
                6: ('ftr', False),
                7: ('orb_pct', False),
                8: ('drb_pct', False),
                9: ('ast_pct', False),
                10: ('tov_pct', True),      # reverse - menos es mejor
                11: ('stl_pct', False),
                12: ('blk_pct', False),
                13: ('usage', False),       # USG%
                14: ('orating', False),     # ORtg
                15: ('drating', True),      # DRtg (reverse)
                16: ('val_pg', False)
            }

            # Calcular cuartiles para cada campo
            quartiles = {}
            for col_idx, (field_key, _) in ADVANCED_STAT_FIELDS.items():
                values = []
                for p in all_player_stats:
                    val = p.get(field_key, 0)
                    # Convertir minutes_per_game de segundos a minutos para cuartiles
                    if field_key == 'minutes_per_game' and val:
                        val = val / 60
                    if val and val != 0:
                        values.append(val)

                if len(values) >= 4:
                    quartiles[field_key] = calculate_quartiles(values)

            # Poblar campos avanzados con coloreado
            from .table_items import NumericTableWidgetItem

            for col_idx, (field_name, reverse) in ADVANCED_STAT_FIELDS.items():
                value = player.get(field_name, 0)

            # Formatear valor
                if field_name == 'games_played':
                    formatted_value = str(int(value)) if value else "0"
                elif field_name == 'minutes_per_game':
                    # Convertir de segundos a minutos para mostrar
                    value_in_minutes = value / 60 if value else 0
                    formatted_value = f"{value_in_minutes:.1f}"
                    value = value_in_minutes  # Actualizar para coloreado
                elif field_name in ['ppg', 'val_pg']:
                    formatted_value = f"{value:.1f}" if value else "0.0"
                elif field_name in ['orating', 'drating']:
                    formatted_value = f"{value:.1f}" if value else "0.0"
                elif field_name in ['three_pr', 'ftr']:
                    # Ya vienen como porcentajes desde el calculador
                    formatted_value = f"{value:.1f}%" if value else "0.0%"
                else:
                    formatted_value = f"{value:.1f}%" if value else "0.0%"

                # Crear item
                item = NumericTableWidgetItem(value if value else 0, formatted_value)

                # Aplicar coloreado por cuartiles
                if value and value != 0 and field_name in quartiles:
                    color = get_quartile_color(value, quartiles[field_name], reverse)
                    item.setBackground(color)

                table.setItem(row, col_idx, item)

            # Solo ajustar altura de filas, mantener anchos fijos
            table.resizeRowsToContents()

            # Calcular tamaño total necesario con margen de seguridad
            total_width = sum(table.columnWidth(i) for i in range(table.columnCount())) + 4
            total_height = (sum(table.rowHeight(i) for i in range(table.rowCount())) +
                          table.horizontalHeader().height() + 4)

            # Establecer tamaño mínimo (no fijo) para permitir expansión si es necesario
            table.setMinimumSize(total_width, total_height)
            table.resize(total_width, total_height)

            # Evitar que la ventana se muestre en pantalla (prevenir parpadeo)
            table.setAttribute(Qt.WidgetAttribute.WA_DontShowOnScreen, True)

            # Renderizar y capturar
            table.show()
            app.processEvents()
            
            # Forzar otro ajuste después de mostrar
            app.processEvents()
            
            # Capturar usando el rect completo
            pixmap = table.grab(table.rect())
            table.hide()

            # Guardar
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            temp_file.close()

            if pixmap.save(temp_file.name, "PNG", quality=100):
                return temp_file.name
            else:
                return None

        except Exception as e:
            import traceback
            traceback.print_exc()
            return None

    def _generate_heatmap(self, shots: List[Dict], player_name: str) -> Optional[str]:
        """
        Genera un mapa de calor de tiros y lo guarda en un archivo temporal.

        Args:
            shots: Lista de tiros
            player_name: Nombre de la jugadora

        Returns:
            Ruta del archivo temporal con el gráfico
        """
        try:
            # Filtrar solo tiros anotados (made shots)
            made_shots = [s for s in shots if int(s.get('m', 0)) == 1]

            # Calcular estadísticas
            made_count = len(made_shots)
            total_count = len(shots)
            accuracy = (made_count / total_count * 100) if total_count > 0 else 0

            # Verificar que hay suficientes tiros para generar heatmap (mínimo 3 puntos)
            if made_count < 3:
                return None  # No generar gráfico si hay muy pocos aciertos

            # Generar heatmap solo con tiros anotados
            fig = self.shot_visualizer.plot_heatmap(
                shots=made_shots,
                title=f"{player_name} - Aciertos\n{made_count}/{total_count} ({accuracy:.1f}%)",
                figsize=(8, 8),
                alpha=0.6
            )

            # Guardar en archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            fig.savefig(temp_file.name, dpi=150, bbox_inches='tight')
            plt.close(fig)

            return temp_file.name

        except Exception as e:
            import traceback
            traceback.print_exc()
            return None

    def _generate_zone_chart(self, shots: List[Dict], player_name: str) -> Optional[str]:
        """
        Genera un gráfico de tiros por zonas y lo guarda en un archivo temporal.

        Args:
            shots: Lista de tiros
            player_name: Nombre de la jugadora

        Returns:
            Ruta del archivo temporal con el gráfico
        """
        try:
            # Convertir shots para análisis de zonas (igual que en ShotChartWindow)
            from shotcharts.coordinate_utils import convert_shots_for_zone_analysis
            processed_shots = convert_shots_for_zone_analysis(shots)

            # Analizar rendimiento por zonas
            stats = self.zone_analyzer.analyze_zone_performance(processed_shots)

            # Calcular estadísticas
            made_count = sum(1 for s in shots if int(s.get('m', 0)) == 1)
            total_count = len(shots)
            accuracy = (made_count / total_count * 100) if total_count > 0 else 0

            # Crear visualización de zonas usando ZoneAnalyzer
            fig = self.zone_analyzer.plot_zone_analysis(
                stats=stats,
                title=f"{player_name} - Análisis por Zonas\n{made_count}/{total_count} ({accuracy:.1f}%)",
                figsize=(8, 8)
            )

            # Guardar en archivo temporal
            temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
            fig.savefig(temp_file.name, dpi=150, bbox_inches='tight')
            plt.close(fig)

            return temp_file.name

        except Exception as e:
            import traceback
            traceback.print_exc()
            return None

    def _generate_radar_chart(
        self,
        player: Dict,
        all_player_stats: List[Dict]
    ) -> Optional[str]:
        """
        Genera un radar chart para la jugadora usando EXACTAMENTE la misma lógica que RadarChartWindow.

        Args:
            player: Estadísticas de la jugadora
            all_player_stats: Estadísticas de todas las jugadoras (contexto)

        Returns:
            Ruta del archivo temporal con el gráfico
        """
        try:
            # Importar RadarChartWindow para usar su lógica exacta
            from .radar_window import RadarChartWindow
            from PyQt6.QtWidgets import QApplication

            # Asegurar que existe QApplication
            app = QApplication.instance()
            if app is None:
                app = QApplication([])

            # Crear ventana de radar (sin mostrarla)
            radar_window = RadarChartWindow(all_player_stats, player, parent=None)

            # Actualizar el chart (genera la figura internamente)
            radar_window.update_chart()

            # Obtener la figura del canvas
            if radar_window.canvas and radar_window.canvas.figure:
                fig = radar_window.canvas.figure

                # Guardar en archivo temporal
                temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.png')
                fig.savefig(temp_file.name, dpi=150, bbox_inches='tight')

                # Cerrar la figura para liberar memoria
                plt.close(fig)

                # Cerrar ventana sin mostrar
                radar_window.close()
                radar_window.deleteLater()

                return temp_file.name
            else:
                return None

        except Exception as e:
            import traceback
            traceback.print_exc()
            return None

