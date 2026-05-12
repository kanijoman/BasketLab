"""DOCX rendering helpers for individual scouting reports.

Extracted from individual_scouting_service.py to keep that module under 500 lines.
All functions here depend only on python-docx and numpy — no FastAPI, no DB.
"""
from __future__ import annotations

from typing import Any, Dict, List, Optional


def write_table(
    doc: Any,
    headers: List[str],
    rows: List[List[str]],
    col_widths_cm: Optional[List[float]] = None,
    no_wrap: bool = False,
) -> None:
    """Compact table with bold header and small font. Optional per-column widths in cm."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "auto")
        tblW.set(qn("w:w"), "0")

    def _apply_cell(cell, is_header: bool, width_cm: Optional[float]) -> None:
        if width_cm:
            dxa = int(width_cm * 567)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(dxa))
        if no_wrap:
            tcPr = cell._tc.get_or_add_tcPr()
            nw = OxmlElement("w:noWrap")
            tcPr.append(nw)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = is_header
            run.font.size = Pt(6.5 if is_header else 7)

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _apply_cell(cell, True, col_widths_cm[i] if col_widths_cm and i < len(col_widths_cm) else None)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            _apply_cell(cell, False, col_widths_cm[c_idx] if col_widths_cm and c_idx < len(col_widths_cm) else None)


def markdown_to_docx(doc: Any, text: str) -> None:
    """Convert simple markdown (**bold**, bullets) into docx paragraphs."""
    import re

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            line = line[2:]
        else:
            p = doc.add_paragraph()
        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
            run = p.add_run(part)
            if i % 2 == 1:
                run.font.bold = True


def set_cell_shading(cell: Any, hex_fill: str) -> None:
    """Apply background fill colour to a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def quartile_fill(
    value: Optional[float],
    all_vals: List[Optional[float]],
    reverse: bool = False,
) -> Optional[str]:
    """Q4=green C6EFCE, Q3=yellow FFEB9C, Q2=orange FFD9B3, Q1=red FFC7CE."""
    import numpy as np
    vals = [float(v) for v in all_vals if v is not None]
    if len(vals) < 4 or value is None:
        return None
    q25, q50, q75 = (float(np.percentile(vals, p)) for p in (25, 50, 75))
    v = float(value)
    if not reverse:
        if v >= q75: return "C6EFCE"
        if v >= q50: return "FFEB9C"
        if v >= q25: return "FFD9B3"
        return "FFC7CE"
    else:
        if v <= q25: return "C6EFCE"
        if v <= q50: return "FFEB9C"
        if v <= q75: return "FFD9B3"
        return "FFC7CE"


def write_player_stats_table(
    doc: Any,
    headers: List[str],
    formatted: List[str],
    raw: List[Optional[float]],
    all_raw: List[List[Optional[float]]],
    reverse_cols: Optional[set] = None,
) -> None:
    """Header row + single data row with per-cell quartile background coloring."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rev = reverse_cols or set()
    t = doc.add_table(rows=2, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(7)
    for i, val in enumerate(formatted):
        cell = t.rows[1].cells[i]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
        if i < len(raw) and i < len(all_raw):
            fill = quartile_fill(raw[i], all_raw[i], reverse=i in rev)
            if fill:
                set_cell_shading(cell, fill)
