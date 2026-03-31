"""Individual team scouting DOCX builder (web tier). Returns raw bytes for FastAPI download.

Layout per player: header → identity → avg stats → totals → advanced stats →
shot profile → radar → AI notes → page break. FEB only for shot data.
"""
from __future__ import annotations

import io
import logging
import tempfile
from datetime import datetime
from typing import Any, Dict, List, Optional

import requests

logger = logging.getLogger(__name__)

_TEAM_LOGO = "https://imagenes.feb.es/imagen.aspx?i={tid}&ti=1"

# ---------------------------------------------------------------------------
# Module-level helpers
# ---------------------------------------------------------------------------

def _safe(val: Any, decimals: int = 1, suffix: str = "") -> str:
    if val is None:
        return "-"
    try:
        return f"{float(val):.{decimals}f}{suffix}"
    except (TypeError, ValueError):
        return str(val)


def _fmt_min(val: Any) -> str:
    """Format float minutes (e.g. 350.5) as MM:SS string."""
    if val is None:
        return "-"
    try:
        m = float(val)
        return f"{int(m)}:{int(round((m % 1) * 60)):02d}"
    except (TypeError, ValueError):
        return str(val)

def _fetch_bytes(url: str) -> Optional[bytes]:
    try:
        r = requests.get(url, timeout=6)
        r.raise_for_status()
        return r.content
    except Exception:
        return None


def _extract_shots(coll: Any, player_id: str) -> List[Dict]:
    """Return FIBA-coordinate shots for one player from a FEB MongoDB collection."""
    from src.shotcharts.coordinate_utils import convert_feb_to_fiba

    shots: List[Dict] = []
    cursor = coll.find({"SHOTCHART.SHOTS": {"$exists": True}}, {"SHOTCHART": 1})
    for doc in cursor:
        sc = doc.get("SHOTCHART", {})
        # Build (team_idx, dorsal) → player_id map for this game
        dorsal_map: Dict = {}
        for ti, team_entry in enumerate(sc.get("TEAM", [])):
            for pl in team_entry.get("PLAYER", []):
                dorsal_map[(ti, str(pl.get("no", "")))] = str(pl.get("id", ""))
        for s in sc.get("SHOTS", []):
            try:
                ti = int(s.get("team"))
            except (TypeError, ValueError):
                continue
            if dorsal_map.get((ti, str(s.get("player", "")))) != player_id:
                continue
            fx, fy = float(s.get("x", 0)), float(s.get("y", 0))
            m_val = int(s.get("m", 0))
            x, y = convert_feb_to_fiba(fx, fy, ti)
            shots.append({
                # FIBA coords + bool — consumed by ZoneAnalyzer
                "x": x, "y": y, "made": m_val == 1,
                # FEB coords + int fields — consumed by ShotChartVisualizer
                "x_feb": fx, "y_feb": fy,
                "m": m_val,
                "team": ti,
                "player": str(s.get("player", "")),
                "quarter": int(s.get("quarter", 1) or 1),
            })
    return shots


def _fig_to_buf(fig: Any) -> io.BytesIO:
    buf = io.BytesIO()
    fig.savefig(buf, format="png", bbox_inches="tight", dpi=100)
    buf.seek(0)
    return buf


def _write_table(
    doc: Any,
    headers: List[str],
    rows: List[List[str]],
    col_widths_cm: Optional[List[float]] = None,
    no_wrap: bool = False,
) -> None:
    """Compact table with bold header and small font. Optional per-column widths in cm."""
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.autofit = False
    tbl = t._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is not None:
        tblW = tblPr.find(qn("w:tblW"))
        if tblW is None:
            tblW = OxmlElement("w:tblW")
            tblPr.append(tblW)
        tblW.set(qn("w:type"), "auto")
        tblW.set(qn("w:w"), "0")

    def _apply_cell(cell, is_header: bool, width_cm: Optional[float]) -> None:
        if width_cm:
            # Force exact column width via XML — python-docx cell.width is only advisory
            dxa = int(width_cm * 567)  # 1 cm = 567 dxa (twentieths of a point)
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:type"), "dxa")
            tcW.set(qn("w:w"), str(dxa))
        if no_wrap:
            tcPr = cell._tc.get_or_add_tcPr()
            nw = OxmlElement("w:noWrap")
            tcPr.append(nw)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = is_header
            run.font.size = Pt(6.5 if is_header else 7)

    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        _apply_cell(cell, True, col_widths_cm[i] if col_widths_cm and i < len(col_widths_cm) else None)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = t.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            _apply_cell(cell, False, col_widths_cm[c_idx] if col_widths_cm and c_idx < len(col_widths_cm) else None)


def _markdown_to_docx(doc: Any, text: str) -> None:
    """Convert simple markdown (**bold**, bullets) into docx paragraphs."""
    import re

    for line in text.splitlines():
        line = line.strip()
        if not line:
            continue
        if line.startswith(("- ", "• ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            line = line[2:]
        else:
            p = doc.add_paragraph()
        for i, part in enumerate(re.split(r"\*\*(.+?)\*\*", line)):
            run = p.add_run(part)
            if i % 2 == 1:
                run.font.bold = True


def _set_cell_shading(cell: Any, hex_fill: str) -> None:
    """Apply background fill colour to a python-docx table cell."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_fill)
    tcPr.append(shd)


def _quartile_fill(
    value: Optional[float],
    all_vals: List[Optional[float]],
    reverse: bool = False,
) -> Optional[str]:
    """Q4=green C6EFCE, Q3=yellow FFEB9C, Q2=orange FFD9B3, Q1=red FFC7CE (inverted when reverse)."""
    import numpy as np
    vals = [float(v) for v in all_vals if v is not None]
    if len(vals) < 4 or value is None:
        return None
    q25, q50, q75 = (float(np.percentile(vals, p)) for p in (25, 50, 75))
    v = float(value)
    if not reverse:
        if v >= q75: return "C6EFCE"
        if v >= q50: return "FFEB9C"
        if v >= q25: return "FFD9B3"
        return "FFC7CE"
    else:
        if v <= q25: return "C6EFCE"
        if v <= q50: return "FFEB9C"
        if v <= q75: return "FFD9B3"
        return "FFC7CE"


def _write_player_stats_table(
    doc: Any,
    headers: List[str],
    formatted: List[str],
    raw: List[Optional[float]],
    all_raw: List[List[Optional[float]]],
    reverse_cols: Optional[set] = None,
) -> None:
    """Header row + single data row with per-cell quartile background coloring."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    rev = reverse_cols or set()
    t = doc.add_table(rows=2, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.bold = True
            run.font.size = Pt(7)
    for i, val in enumerate(formatted):
        cell = t.rows[1].cells[i]
        cell.text = val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(8)
        if i < len(raw) and i < len(all_raw):
            fill = _quartile_fill(raw[i], all_raw[i], reverse=i in rev)
            if fill:
                _set_cell_shading(cell, fill)


def _prep_player_for_radar(p: Dict) -> Dict:
    """Add field aliases expected by RadarChart.calculate_metrics_from_stats() (Qt radar_window.py pattern)."""
    d = dict(p)
    d.setdefault("ts",      d.get("true_shooting", 0))
    d.setdefault("usage",   d.get("usage_pct", 0))
    d.setdefault("tov_pct", d.get("tov_pct_adv", 0))
    return d


# ---------------------------------------------------------------------------
# Main builder class
# ---------------------------------------------------------------------------

class IndividualScoutingDocxBuilder:
    """Build a team scouting DOCX with one full page per player.

    Usage::

        builder = IndividualScoutingDocxBuilder(
            collection="FEB_LF2_2025_A",
            team_name="Club Baloncesto",
            db=handler,
        )
        docx_bytes = builder.build()
    """

    def __init__(
        self,
        collection: str,
        team_name: str,
        db: Any,
        *,
        include_ai_notes: bool = True,
        provider: str = "groq",
    ) -> None:
        self.collection = collection
        self.team_name = team_name
        self.db = db
        self.include_ai_notes = include_ai_notes
        self.provider = provider
        self.is_fbcyl = "FBCYL" in collection.upper()

    def build(self) -> bytes:
        """Build the DOCX and return raw bytes.  Returns b'' if no players found."""
        from docx import Document
        from docx.shared import Cm
        from src.services.player_stats_service import PlayerStatsService
        from src.ui.player_data_fetcher import PlayerDataFetcher

        svc = PlayerStatsService(self.db)
        all_players = svc.load_season_data(self.collection)
        team_players = [p for p in all_players if p.get("team_name") == self.team_name]
        if not team_players:
            return b""

        fetcher = PlayerDataFetcher(self.db, self.collection)

        # Team logo (resolved once from first player's team_id)
        _, _, team_id = fetcher.get_player_dorsal_and_photo(
            str(team_players[0].get("player_id", ""))
        )
        logo_bytes = _fetch_bytes(_TEAM_LOGO.format(tid=team_id)) if team_id else None

        # MongoDB collection for shot data (FEB only)
        coll = None
        if not self.is_fbcyl:
            try:
                coll = self.db.connection.get_collection(self.collection)
            except Exception:
                pass

        doc = Document()
        for section in doc.sections:
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(1.8)
            section.right_margin = Cm(1.8)

        for i, player in enumerate(team_players):
            self._add_player_page(doc, player, all_players, fetcher, coll, logo_bytes, team_id)
            if i < len(team_players) - 1:
                doc.add_page_break()

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ------------------------------------------------------------------
    # Per-player sections
    # ------------------------------------------------------------------

    def _add_player_page(self, doc, player, all_players, fetcher, coll, logo_bytes, team_id):
        name = player.get("player_name", "—")
        pid = str(player.get("player_id", ""))
        self._add_doc_header(doc, logo_bytes)
        self._add_identity_block(doc, player, fetcher, pid, team_id)
        self._add_basic_stats_tables(doc, player, all_players)
        self._add_advanced_stats_table(doc, player, all_players)
        if coll is not None:
            shots = _extract_shots(coll, pid)
            self._add_shot_profile(doc, shots, name)
        self._add_radar_chart(doc, player, all_players, name)
        self._add_ai_notes(doc, player)

    def _add_doc_header(self, doc, logo_bytes):
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        left, right = t.rows[0].cells[0], t.rows[0].cells[1]

        if logo_bytes:
            try:
                from PIL import Image
                img = Image.open(io.BytesIO(logo_bytes))
                with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                    img.save(f.name)
                    lp = left.paragraphs[0]
                    lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    lp.add_run().add_picture(f.name, height=Inches(0.6))
            except Exception:
                left.paragraphs[0].text = self.team_name
        else:
            left.paragraphs[0].text = self.team_name

        p = right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run(f"Scouting Individual — {self.team_name}")
        r.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)
        dp = right.add_paragraph()
        dp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        dp.add_run(datetime.now().strftime("%d/%m/%Y")).font.size = Pt(9)

    def _add_identity_block(self, doc, player, fetcher, pid, team_id):
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT

        dorsal, photo_url, _ = fetcher.get_player_dorsal_and_photo(pid)
        birth_date, age, height = fetcher.get_player_birth_info(pid, team_id)

        t = doc.add_table(rows=1, cols=2)
        left, right = t.rows[0].cells[0], t.rows[0].cells[1]

        p = left.paragraphs[0]
        r = p.add_run(f"#{dorsal}  {player.get('player_name', '—')}")
        r.bold = True
        r.font.size = Pt(14)

        for label, val in [
            ("Nac.", birth_date or "___________"),
            ("Edad", f"{age} años" if age else "___________"),
            ("Altura", height or "___________"),
            ("Posición", "___________"),
        ]:
            lp = left.add_paragraph()
            lp.add_run(f"{label}: ").bold = True
            lp.add_run(str(val))

        right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        photo_added = False
        if photo_url:
            try:
                pb = _fetch_bytes(photo_url)
                if pb:
                    from PIL import Image
                    img = Image.open(io.BytesIO(pb)).convert("RGB")
                    with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as f:
                        img.save(f.name)
                        rp = right.paragraphs[0]
                        rp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        rp.add_run().add_picture(f.name, height=Inches(1.3))
                    photo_added = True
            except Exception:
                pass
        if not photo_added:
            right.paragraphs[0].text = "[FOTO]"

    def _add_basic_stats_tables(self, doc, player, all_players):
        doc.add_heading('1. Estadística General', level=1)

        _AVG_DEFS = [
            ("PJ",  "games_played",                   False),
            ("MIN", "minutes_per_game",               False),
            ("PTS", "points_per_game",                False),
            ("TL%", "fg1_percentage",                 False),
            ("T2%", "fg2_percentage",                 False),
            ("T3%", "fg3_percentage",                 False),
            ("RO",  "offensive_rebounds_per_game",    False),
            ("RD",  "defensive_rebounds_per_game",    False),
            ("RT",  "rebounds_per_game",              False),
            ("AST", "assists_per_game",               False),
            ("ROB", "steals_per_game",                False),
            ("BP",  "turnovers_per_game",             True),
            ("TAP", "blocks_per_game",                False),
            ("FP",  "fouls_per_game",                 True),
            ("VAL", "valoracion_per_game",            False),
        ]
        _pct = {"fg1_percentage", "fg2_percentage", "fg3_percentage"}
        _min = {"minutes_per_game"}
        headers = [d[0] for d in _AVG_DEFS]
        fmt = [
            _fmt_min(player.get(d[1])) if d[1] in _min
            else _safe(player.get(d[1]), suffix="%") if d[1] in _pct
            else _safe(player.get(d[1]))
            for d in _AVG_DEFS
        ]
        raw = [player.get(d[1]) for d in _AVG_DEFS]
        all_raw = [[p.get(d[1]) for p in all_players] for d in _AVG_DEFS]
        rev = {i for i, d in enumerate(_AVG_DEFS) if d[2]}

        doc.add_paragraph().add_run("Promedios por partido:").bold = True
        _write_player_stats_table(doc, headers, fmt, raw, all_raw, rev)

        doc.add_paragraph().add_run("Totales acumulados:").bold = True
        tm = player.get("total_minutes", 0) or 0
        # Widths in cm: PJ MIN PTS TL   T2   T3   RO  RD  RT  AST ROB BP  TAP FP  VAL
        _TW = [0.7, 1.5, 0.8, 2.0, 2.0, 2.0, 0.7, 0.7, 0.7, 0.8, 0.8, 0.8, 0.8, 0.8, 0.8]
        # \u2060 = Unicode word-joiner: prevents Word from breaking at the "/" character
        _s = "\u2060/\u2060"
        _write_table(doc, [
            "PJ", "MIN", "PTS", "TL", "T2", "T3",
            "RO", "RD", "RT", "AST", "ROB", "BP", "TAP", "FP", "VAL",
        ], [[
            str(player.get("games_played", 0)),
            _fmt_min(tm),
            str(player.get("total_pts", 0)),
            f"{player.get('total_p1m', 0)}{_s}{player.get('total_p1a', 0)}",
            f"{player.get('total_p2m', 0)}{_s}{player.get('total_p2a', 0)}",
            f"{player.get('total_p3m', 0)}{_s}{player.get('total_p3a', 0)}",
            str(player.get("total_ro", 0)),
            str(player.get("total_rd", 0)),
            str(player.get("total_rt", 0)),
            str(player.get("total_assist", 0)),
            str(player.get("total_st", 0)),
            str(player.get("total_to", 0)),
            str(player.get("total_bs", 0)),
            str(player.get("total_pf", 0)),
            str(player.get("total_val", 0)),
        ]], col_widths_cm=_TW, no_wrap=True)

    def _add_advanced_stats_table(self, doc, player, all_players):
        doc.add_heading('2. Estadística Avanzada', level=1)

        _ADV_DEFS = [
            ("TS%",    "true_shooting",    False),
            ("eFG%",   "efg_percentage",   False),
            ("3PAr",   "three_point_rate", False),
            ("FTr",    "free_throw_rate",  False),
            ("ORB%",   "orb_pct",          False),
            ("DRB%",   "drb_pct",          False),
            ("AST%",   "ast_pct",          False),
            ("TO%",    "tov_pct_adv",      True),
            ("ROB%",   "stl_pct",          False),
            ("TAP%",   "blk_pct",          False),
            ("USG%",   "usage_pct",        False),
            ("ORtg",   "orating",          False),
            ("DRtg",   "drating",          True),
            ("NetRtg", "net_rtg",          False),
        ]
        _pct = {
            "true_shooting", "efg_percentage", "three_point_rate", "free_throw_rate",
            "orb_pct", "drb_pct", "ast_pct", "tov_pct_adv", "stl_pct", "blk_pct", "usage_pct",
        }
        headers = [d[0] for d in _ADV_DEFS]
        fmt = [_safe(player.get(d[1]), suffix="%") if d[1] in _pct else _safe(player.get(d[1])) for d in _ADV_DEFS]
        raw = [player.get(d[1]) for d in _ADV_DEFS]
        all_raw = [[p.get(d[1]) for p in all_players] for d in _ADV_DEFS]
        rev = {i for i, d in enumerate(_ADV_DEFS) if d[2]}
        _write_player_stats_table(doc, headers, fmt, raw, all_raw, rev)

    def _add_shot_profile(self, doc, shots: List[Dict], player_name: str) -> None:
        if not shots:
            return
        doc.add_heading('3. Perfil de Lanzamiento', level=1)
        import matplotlib.pyplot as plt
        from docx.shared import Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from src.shotcharts.zone_analysis import ZoneAnalyzer
        from src.shotcharts.shot_visualizer import ShotChartVisualizer

        row_t = doc.add_table(rows=1, cols=2)
        # Heatmap (left)
        viz_shots = [
            {"x": s["x_feb"], "y": s["y_feb"], "m": s["m"],
             "team": s["team"], "player": s["player"], "quarter": s["quarter"]}
            for s in shots
        ]
        hfig = None
        try:
            hfig = ShotChartVisualizer().plot_heatmap(viz_shots, title=player_name, figsize=(12, 6))
        except Exception:
            try:
                hfig = ShotChartVisualizer().plot_shots(viz_shots, title=player_name, figsize=(12, 6))
            except Exception as exc:
                logger.warning("Heatmap/shot chart failed for %s: %s", player_name, exc)
        if hfig is not None:
            try:
                hbuf = _fig_to_buf(hfig)
                plt.close(hfig)
                hp = row_t.rows[0].cells[0].paragraphs[0]
                hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                hp.add_run().add_picture(hbuf, height=Inches(2.5))
            except Exception as exc:
                logger.warning("Embedding heatmap failed: %s", exc)
        # Zone chart (right)
        try:
            az = ZoneAnalyzer()
            zstats = az.analyze_zone_performance(shots)
            zfig = az.plot_zone_analysis(zstats, title=f"{player_name} — Zonas", figsize=(12, 6))
            zbuf = _fig_to_buf(zfig)
            plt.close(zfig)
            zp = row_t.rows[0].cells[1].paragraphs[0]
            zp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            zp.add_run().add_picture(zbuf, height=Inches(2.5))
        except Exception as exc:
            logger.warning("Zone chart failed for %s: %s", player_name, exc)

    def _add_radar_chart(self, doc, player, all_players, player_name: str) -> None:
        try:
            import matplotlib.pyplot as plt
            from docx.shared import Inches
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from src.visualization.radar_chart import RadarChart

            h = doc.add_heading('4. Perfil de Juego', level=1)
            h.paragraph_format.page_break_before = True
            radar = RadarChart(figsize=(8, 6))
            player_metrics = radar.calculate_metrics_from_stats(_prep_player_for_radar(player))
            league_metrics = [radar.calculate_metrics_from_stats(_prep_player_for_radar(p)) for p in all_players]
            fig = radar.create_chart(player_metrics, league_metrics, player_name)
            buf = _fig_to_buf(fig)
            plt.close(fig)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(buf, width=Inches(5.5))
        except Exception as exc:
            logger.warning("Radar chart generation failed: %s", exc)

    def _add_ai_notes(self, doc, player: Dict) -> None:
        doc.add_heading('5. Notas del cuerpo técnico', level=1)
        notes = self._call_ai(player) if self.include_ai_notes else None
        if notes:
            _markdown_to_docx(doc, notes)
        else:
            for _ in range(6):
                doc.add_paragraph("_" * 80)

    def _call_ai(self, player: Dict) -> Optional[str]:
        """Synchronous Groq call for brief player scouting notes."""
        try:
            import openai as _oai
            from src.ai.prompts import PROMPT_PLAYER_NOTES_BRIEF
            from src.ai.config import AnalysisConfig

            AnalysisConfig.load_api_keys()
            if not AnalysisConfig.GROQ_API_KEY:
                return None

            ctx = (
                f"{player.get('player_name', '')} ({player.get('team_name', '')})\n"
                f"PJ={player.get('games_played', 0)} "
                f"MIN={_safe(player.get('minutes_per_game'))} "
                f"PTS={_safe(player.get('points_per_game'))} "
                f"REB={_safe(player.get('rebounds_per_game'))} "
                f"AST={_safe(player.get('assists_per_game'))}\n"
                f"T2={_safe(player.get('fg2_percentage'))}% "
                f"T3={_safe(player.get('fg3_percentage'))}% "
                f"TL={_safe(player.get('fg1_percentage'))}%\n"
                f"TS={_safe(player.get('true_shooting'))}% "
                f"eFG={_safe(player.get('efg_percentage'))}% "
                f"USG={_safe(player.get('usage_pct'))}%\n"
                f"ORtg={_safe(player.get('orating'))} "
                f"DRtg={_safe(player.get('drating'))} "
                f"NetRtg={_safe(player.get('net_rtg'))}"
            )
            client = _oai.OpenAI(
                api_key=AnalysisConfig.GROQ_API_KEY,
                base_url="https://api.groq.com/openai/v1",
            )
            resp = client.chat.completions.create(
                model=AnalysisConfig.GROQ_MODELS.get("fast", "llama-3.3-70b-versatile"),
                messages=[
                    {"role": "system", "content": PROMPT_PLAYER_NOTES_BRIEF},
                    {"role": "user", "content": ctx},
                ],
                temperature=0.7,
                max_tokens=400,
            )
            return resp.choices[0].message.content.strip()
        except Exception as exc:
            logger.warning("AI notes generation failed: %s", exc)
            return None
