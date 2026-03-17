"""
Report generation service — Phase 5.

Builds PDF (fpdf2) and DOCX (python-docx) reports from web-layer data.
All methods return raw bytes so the router can stream them directly.
"""
from __future__ import annotations

import io
from datetime import datetime
from typing import Any

from fpdf import FPDF

from src.services.player_stats_service import PlayerStatsService
from src.services.team_stats_service import TeamStatsService

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

_NOW = lambda: datetime.now().strftime("%d/%m/%Y")

_STAT_LABELS = [
    ("points_per_game",   "Puntos"),
    ("rebounds_per_game", "Rebotes"),
    ("assists_per_game",  "Asistencias"),
    ("steals_per_game",   "Robos"),
    ("blocks_per_game",   "Tapones"),
    ("minutes_per_game",  "Min/partido"),
    ("games_played",      "Partidos"),
]


def _safe(val: Any, decimals: int = 1) -> str:
    if val is None:
        return "-"
    try:
        return f"{float(val):.{decimals}f}"
    except (TypeError, ValueError):
        return _cp1252(str(val))


def _cp1252(s: str) -> str:
    """Sanitize a string so FPDF2 core (Helvetica/CP1252) fonts don't raise.

    Characters outside Windows-1252 are replaced with '?'.
    """
    return s.encode("cp1252", errors="replace").decode("cp1252")


# ---------------------------------------------------------------------------
# Base PDF class
# ---------------------------------------------------------------------------

class _BasePDF(FPDF):
    """FPDF subclass with shared header/footer for BasketLab reports."""

    title_text: str = "BasketLab"

    def header(self) -> None:
        self.set_font("Helvetica", "B", 14)
        self.set_fill_color(30, 30, 30)
        self.set_text_color(255, 255, 255)
        self.cell(0, 10, self.title_text, align="C", fill=True, new_x="LMARGIN", new_y="NEXT")
        self.ln(2)

    def footer(self) -> None:
        self.set_y(-12)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 6, f"BasketLab | Generado el {_NOW()} | Pag. {self.page_no()}", align="C")

    def section_title(self, text: str) -> None:
        self.set_font("Helvetica", "B", 11)
        self.set_fill_color(50, 50, 80)
        self.set_text_color(255, 255, 255)
        self.cell(0, 8, text, fill=True, new_x="LMARGIN", new_y="NEXT")
        self.set_text_color(0, 0, 0)
        self.ln(1)

    def kv_row(self, key: str, value: str, shade: bool = False) -> None:
        if shade:
            self.set_fill_color(240, 240, 245)
        else:
            self.set_fill_color(255, 255, 255)
        self.set_font("Helvetica", "", 10)
        self.cell(70, 7, key, fill=True)
        self.set_font("Helvetica", "B", 10)
        self.cell(0, 7, value, fill=True, new_x="LMARGIN", new_y="NEXT")
        self.set_font("Helvetica", "", 10)


# ---------------------------------------------------------------------------
# Player scouting DOCX
# ---------------------------------------------------------------------------

def build_player_scouting_docx(collection: str, player_id: str, db: Any) -> bytes:
    """Return DOCX bytes with a one-page scouting sheet for the given player."""
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    svc   = PlayerStatsService(db)
    rows  = svc.load_season_data(collection)
    player = next((r for r in rows if str(r.get("player_id")) == str(player_id)), None)

    doc = Document()
    # Title
    h  = doc.add_heading("", level=0)
    run = h.add_run("BasketLab — Informe de Scouting Individual")
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x4E)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if player is None:
        doc.add_paragraph(f"No se encontraron datos para el jugador ID {player_id}.")
    else:
        # Identity
        doc.add_heading(player.get("player_name", "—"), level=1)
        info_para = doc.add_paragraph()
        info_para.add_run(f"Equipo: ").bold = True
        info_para.add_run(player.get("team_name", "—"))
        info_para.add_run("   Colección: ").bold = True
        info_para.add_run(collection)
        info_para.add_run(f"   Fecha: ").bold = True
        info_para.add_run(_NOW())

        # Stats table
        doc.add_heading("Estadísticas por partido", level=2)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Estadística"
        hdr[1].text = "Valor"
        for key, label in _STAT_LABELS:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = _safe(player.get(key))

        # Notes section
        doc.add_heading("Notas del cuerpo técnico", level=2)
        for _ in range(8):
            para = doc.add_paragraph("_" * 80)
            para.paragraph_format.space_after = Pt(2)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Team scouting PDF
# ---------------------------------------------------------------------------

def build_team_scouting_pdf(collection: str, team_name: str, db: Any) -> bytes:
    """Return PDF bytes with a team scouting report."""
    svc  = TeamStatsService(db)
    team = svc.get_team_detailed_stats(collection, team_name)
    opp  = svc.get_opponent_detailed_stats(collection, team_name)

    pdf = _BasePDF()
    pdf.title_text = _cp1252(f"Scouting: {team_name}")
    pdf.add_page()

    pdf.section_title(f"Estadisticas ofensivas - {team_name}")
    off_keys = [
        ("points", "Puntos por partido"),
        ("field_goals_attempted", "Tiros intentados"),
        ("field_goals_made", "Tiros anotados"),
        ("three_points_attempted", "3PT intentados"),
        ("three_points_made", "3PT anotados"),
        ("free_throws_attempted", "TL intentados"),
        ("free_throws_made", "TL anotados"),
        ("offensive_rebounds", "Rebotes ofensivos"),
        ("assists",   "Asistencias"),
        ("turnovers", "Pérdidas"),
    ]
    for i, (key, label) in enumerate(off_keys):
        pdf.kv_row(label, _safe(team.get(key)), shade=bool(i % 2))

    pdf.ln(3)
    pdf.section_title(f"Estadisticas defensivas - {team_name}")
    def_keys = [
        ("points", "Puntos rival"),
        ("field_goals_attempted", "Tiros rival intentados"),
        ("field_goals_made", "Tiros rival anotados"),
        ("defensive_rebounds", "Rebotes defensivos"),
        ("steals", "Robos"),
        ("blocks", "Tapones"),
    ]
    for i, (key, label) in enumerate(def_keys):
        pdf.kv_row(label, _safe(opp.get(key)), shade=bool(i % 2))

    return bytes(pdf.output())


# ---------------------------------------------------------------------------
# Season summary PDF
# ---------------------------------------------------------------------------

def build_season_summary_pdf(collection: str, db: Any) -> bytes:
    """Return PDF bytes with a league-wide season summary ranking table."""
    svc   = TeamStatsService(db)
    stats = svc.get_possession_stats(collection)  # list of team dicts w/ pace/oer/der/net_rating

    pdf = _BasePDF()
    pdf.title_text = _cp1252(f"Resumen de temporada - {collection}")
    pdf.add_page()
    pdf.section_title(_cp1252(f"Clasificacion de eficiencia (por posesiones)"))

    # Table header
    cols   = ["Equipo", "Pace", "OER", "DER", "Net Rtg"]
    widths = [70, 28, 28, 28, 28]

    pdf.set_font("Helvetica", "B", 9)
    pdf.set_fill_color(50, 50, 80)
    pdf.set_text_color(255, 255, 255)
    for w, c in zip(widths, cols):
        pdf.cell(w, 8, c, border=1, fill=True, align="C")
    pdf.ln()

    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Helvetica", "", 9)
    for i, row in enumerate(sorted(stats, key=lambda r: r.get("net_rating") or 0, reverse=True)):
        fill = bool(i % 2)
        pdf.set_fill_color(240, 240, 245) if fill else pdf.set_fill_color(255, 255, 255)
        nr = row.get("net_rating", 0) or 0
        pdf.cell(widths[0], 7, _cp1252(str(row.get("team", "-")))[:35],  border=1, fill=fill)
        pdf.cell(widths[1], 7, _safe(row.get("pace")),          border=1, fill=fill, align="C")
        pdf.cell(widths[2], 7, _safe(row.get("oer")),           border=1, fill=fill, align="C")
        pdf.cell(widths[3], 7, _safe(row.get("der")),           border=1, fill=fill, align="C")
        sign = "+" if nr >= 0 else ""
        pdf.cell(widths[4], 7, f"{sign}{nr:.1f}",               border=1, fill=fill, align="C")
        pdf.ln()

    return bytes(pdf.output())
